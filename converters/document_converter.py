import os
import shutil
import subprocess
import tempfile
from pathlib import Path

# Librerie per documenti
from PyPDF2 import PdfReader, PdfWriter, PdfMerger
import docx
try:
    from docx2pdf import convert as docx_to_pdf
except ImportError:
    docx_to_pdf = None

# Fallback imports
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    import textwrap
except ImportError:
    pass

from markdown import markdown
import html2text
from bs4 import BeautifulSoup
import csv
import json
import xml.etree.ElementTree as ET

from .base_converter import BaseConverter

class DocumentConverter(BaseConverter):
    """Convertitore per documenti (PDF, DOC, DOCX, TXT, HTML, MD, ecc.)"""
    
    def __init__(self, source_format=None, target_format=None):
        super().__init__(source_format, target_format)
        # Verifica se Pandoc è disponibile per conversioni avanzate
        self.pandoc_available = self._check_pandoc()
        
    def _check_pandoc(self):
        """Verifica se Pandoc è installato nel sistema"""
        try:
            result = subprocess.run(['pandoc', '--version'], 
                                   stdout=subprocess.PIPE, 
                                   stderr=subprocess.PIPE)
            return result.returncode == 0
        except FileNotFoundError:
            print("Pandoc non è installato. Alcune conversioni potrebbero essere limitate.")
            return False
    
    def get_supported_input_formats(self):
        formats = [
            'pdf', 'doc', 'docx', 'odt', 'txt', 'rtf', 'html', 'htm',
            'md', 'markdown', 'csv', 'json', 'xml'
        ]
        
        # Se Pandoc è disponibile, aggiungi altri formati
        if self.pandoc_available:
            formats.extend(['epub', 'tex', 'org', 'rst', 'adoc', 'pptx', 'xlsx'])
            
        return formats
    
    def get_supported_output_formats(self):
        formats = [
            'pdf', 'docx', 'odt', 'txt', 'rtf', 'html',
            'md', 'markdown', 'csv', 'json'
        ]
        
        # Se Pandoc è disponibile, aggiungi altri formati
        if self.pandoc_available:
            formats.extend(['epub', 'tex', 'org', 'rst', 'adoc'])
            
        return formats
    
    def convert(self, input_path, output_path, **kwargs):
        """
        Converte un documento dal formato sorgente al formato target
        
        Args:
            input_path: percorso del file da convertire
            output_path: percorso dove salvare il file convertito
            **kwargs: parametri aggiuntivi per la conversione:
                - preserve_metadata: bool, mantiene i metadati se possibile
                - paper_size: str, formato carta (A4, Letter, ecc.)
                - margin: dict, margini del documento in mm (top, right, bottom, left)
                - font: str, font da utilizzare
                - font_size: int, dimensione del font
                - header_footer: dict, personalizzazione header/footer
                - toc: bool, genera indice automatico
                - template: str, percorso a un template personalizzato
                - encrypted_pdf: bool, se PDF è crittografato
                - password: str, password per PDF crittografati
        """
        try:
            source_format = os.path.splitext(input_path)[1][1:].lower()
            target_format = os.path.splitext(output_path)[1][1:].lower()
            
            # Estrai opzioni
            preserve_metadata = kwargs.get('preserve_metadata', True)
            paper_size = kwargs.get('paper_size', 'A4')
            margins = kwargs.get('margins', {'top': 20, 'right': 20, 'bottom': 20, 'left': 20})
            font = kwargs.get('font', None)
            font_size = kwargs.get('font_size', None)
            header_footer = kwargs.get('header_footer', None)
            toc = kwargs.get('toc', False)
            template = kwargs.get('template', None)
            encrypted_pdf = kwargs.get('encrypted_pdf', False)
            password = kwargs.get('password', '')
            
            # Per conversioni da/a PDF
            if source_format in ['pdf'] or target_format in ['pdf']:
                return self._convert_pdf(input_path, output_path, source_format, target_format, **kwargs)
            
            # Per conversioni tra formati Office
            elif source_format in ['doc', 'docx', 'odt', 'rtf'] and target_format in ['doc', 'docx', 'odt', 'rtf']:
                return self._convert_office_formats(input_path, output_path, source_format, target_format, **kwargs)
            
            # Per conversioni da/a formati di testo semplice
            elif source_format in ['txt', 'md', 'markdown', 'html', 'htm'] or target_format in ['txt', 'md', 'markdown', 'html', 'htm']:
                return self._convert_text_formats(input_path, output_path, source_format, target_format, **kwargs)
            
            # Per conversioni da/a formati strutturati (CSV, JSON, XML)
            elif source_format in ['csv', 'json', 'xml'] or target_format in ['csv', 'json', 'xml']:
                return self._convert_structured_formats(input_path, output_path, source_format, target_format, **kwargs)
            
            # Utilizzo Pandoc per altre conversioni se disponibile
            elif self.pandoc_available:
                return self._convert_with_pandoc(input_path, output_path, source_format, target_format, **kwargs)
            
            else:
                print(f"Conversione da {source_format} a {target_format} non supportata")
                return False
                
        except Exception as e:
            print(f"Errore durante la conversione del documento: {e}")
            return False

    def _convert_pdf(self, input_path, output_path, source_format, target_format, **kwargs):
        """Gestisce la conversione da/a PDF"""
        
        if source_format == 'pdf' and target_format in ['txt', 'md', 'markdown', 'html']:
            # Conversione da PDF a testo/markdown/html
            try:
                reader = PdfReader(input_path)
                
                # Se il PDF è crittografato, prova a decifrarlo
                if kwargs.get('encrypted_pdf', False) and reader.is_encrypted:
                    if not reader.decrypt(kwargs.get('password', '')):
                        raise ValueError("Password PDF non valida")
                
                # Estrai il testo da tutte le pagine
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n\n"
                
                # Converti in formato appropriato
                if target_format in ['md', 'markdown']:
                    # Tentativo molto semplice di formattazione markdown
                    # Una soluzione più completa richiederebbe l'analisi della struttura del documento
                    text = text.replace('\n\n', '\n\n## ').replace('\t', '    ')
                    text = "# Documento convertito\n\n" + text
                elif target_format == 'html':
                    # Conversione molto semplice, senza analisi della struttura
                    text = f"<html><body><pre>{text}</pre></body></html>"
                
                # Salva il risultato
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
                
                return True
                
            except Exception as e:
                print(f"Errore nella conversione da PDF: {e}")
                return False
                
        elif target_format == 'pdf':
            # Conversione a PDF
            try:
                if source_format == 'docx':
                    # Usa docx2pdf se disponibile (Windows/Office installato)
                    if docx_to_pdf is not None:
                        try:
                            docx_to_pdf(input_path, output_path)
                            return True
                        except Exception as e:
                            print(f"docx2pdf fallito, tento fallback: {e}")
                    
                    # Fallback: Pure Python conversion (Linux/Render)
                    return self._docx_to_pdf_fallback(input_path, output_path)

                elif source_format in ['md', 'markdown', 'txt', 'html', 'htm']:
                    # Usa Pandoc se disponibile
                    if self.pandoc_available:
                        return self._convert_with_pandoc(input_path, output_path, source_format, target_format, **kwargs)
                    else:
                        # Fallback: converti a HTML e poi usa una libreria di rendering HTML a PDF
                        # Questo è un esempio semplificato, una versione completa richiederebbe weasyprint o simili
                        if source_format in ['md', 'markdown']:
                            with open(input_path, 'r', encoding='utf-8') as f:
                                html_content = markdown(f.read())
                        elif source_format == 'txt':
                            with open(input_path, 'r', encoding='utf-8') as f:
                                text = f.read()
                                html_content = f"<html><body><pre>{text}</pre></body></html>"
                        else:  # HTML
                            with open(input_path, 'r', encoding='utf-8') as f:
                                html_content = f.read()
                        
                        # Qui dovresti usare una libreria come weasyprint per convertire HTML a PDF
                        # Questo è solo un segnaposto
                        print("Necessaria una libreria HTML-to-PDF come WeasyPrint per questa conversione")
                        return False
            except Exception as e:
                print(f"Errore nella conversione a PDF: {e}")
                return False
        
        # Se siamo qui, il formato non è supportato
        return self._convert_with_pandoc(input_path, output_path, source_format, target_format, **kwargs)
    
    def _convert_office_formats(self, input_path, output_path, source_format, target_format, **kwargs):
        """Gestisce la conversione tra formati Office (docx, odt, rtf)"""
        
        # Per la maggior parte delle conversioni tra formati Office, Pandoc è l'opzione migliore
        if self.pandoc_available:
            return self._convert_with_pandoc(input_path, output_path, source_format, target_format, **kwargs)
        
        # Implementazione base per docx -> txt
        if source_format == 'docx' and target_format == 'txt':
            try:
                doc = docx.Document(input_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(full_text))
                
                return True
            except Exception as e:
                print(f"Errore nella conversione docx -> txt: {e}")
                return False
        
        # Per altre conversioni, non abbiamo un percorso diretto
        print(f"Conversione diretta da {source_format} a {target_format} non implementata")
        return False
    
    def _convert_text_formats(self, input_path, output_path, source_format, target_format, **kwargs):
        """Gestisce la conversione tra formati di testo (txt, md, html)"""
        try:
            # Carica il contenuto di origine
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Converti in base al formato di destinazione
            if source_format in ['md', 'markdown'] and target_format == 'html':
                # Markdown a HTML
                html_content = markdown(content)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                
            elif source_format in ['html', 'htm'] and target_format in ['md', 'markdown']:
                # HTML a Markdown
                h = html2text.HTML2Text()
                h.ignore_links = False
                md_content = h.handle(content)
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(md_content)
                
            elif source_format in ['html', 'htm', 'md', 'markdown'] and target_format == 'txt':
                # HTML/Markdown a testo semplice
                if source_format in ['html', 'htm']:
                    soup = BeautifulSoup(content, 'html.parser')
                    text_content = soup.get_text()
                else:
                    # Conversione più semplice per markdown -> txt
                    # Rimuove i caratteri speciali di markdown
                    text_content = content.replace('#', '').replace('*', '').replace('_', '')
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                
            elif source_format == 'txt' and target_format in ['html', 'htm']:
                # Testo semplice a HTML
                html_content = f"<html><body><pre>{content}</pre></body></html>"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                
            elif source_format == 'txt' and target_format in ['md', 'markdown']:
                # Testo semplice a Markdown (aggiungendo struttura minima)
                lines = content.split('\n')
                md_content = ""
                for i, line in enumerate(lines):
                    if i == 0:
                        md_content += f"# {line}\n\n"
                    elif line.strip() == '':
                        md_content += "\n"
                    else:
                        md_content += f"{line}\n"
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(md_content)
            
            else:
                # Per altre conversioni, usa Pandoc se disponibile
                if self.pandoc_available:
                    return self._convert_with_pandoc(input_path, output_path, source_format, target_format, **kwargs)
                else:
                    return False
            
            return True
            
        except Exception as e:
            print(f"Errore nella conversione di formati testuali: {e}")
            return False
    
    def _convert_structured_formats(self, input_path, output_path, source_format, target_format, **kwargs):
        """Gestisce la conversione tra formati strutturati (csv, json, xml)"""
        try:
            # CSV -> JSON
            if source_format == 'csv' and target_format == 'json':
                data = []
                with open(input_path, 'r', newline='', encoding='utf-8') as f:
                    reader = csv.DictReader(f)
                    for row in reader:
                        data.append(row)
                
                with open(output_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=2)
                
                return True
            
            # JSON -> CSV
            elif source_format == 'json' and target_format == 'csv':
                with open(input_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # Supporta solo array di oggetti
                if not isinstance(data, list) or not all(isinstance(item, dict) for item in data):
                    raise ValueError("Il JSON deve essere un array di oggetti")
                
                # Estrai tutti i campi possibili
                fieldnames = set()
                for item in data:
                    fieldnames.update(item.keys())
                
                with open(output_path, 'w', newline='', encoding='utf-8') as f:
                    writer = csv.DictWriter(f, fieldnames=sorted(fieldnames))
                    writer.writeheader()
                    writer.writerows(data)
                
                return True
            
            # XML -> JSON
            elif source_format == 'xml' and target_format == 'json':
                def _parse_xml_element(elem):
                    """Converte un elemento XML in dizionario per JSON"""
                    result = {}
                    # Gestisci attributi
                    if elem.attrib:
                        result["@attributes"] = elem.attrib
                    
                    # Gestisci elementi figli
                    children = list(elem)
                    if not children:
                        result["#text"] = elem.text.strip() if elem.text else ""
                    else:
                        child_elements = {}
                        for child in children:
                            tag = child.tag
                            child_dict = _parse_xml_element(child)
                            
                            if tag in child_elements:
                                if not isinstance(child_elements[tag], list):
                                    child_elements[tag] = [child_elements[tag]]
                                child_elements[tag].append(child_dict)
                            else:
                                child_elements[tag] = child_dict
                        
                        result.update(child_elements)
                    
                    return result
                
                # Parse XML
                tree = ET.parse(input_path)
                root = tree.getroot()
                data = {root.tag: _parse_xml_element(root)}
                
                # Save as JSON
                with open(output_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, indent=2)
                
                return True
            
            elif source_format == 'json' and target_format == 'xml':
                # Questa è una conversione semplificata, una versione completa richiederebbe una mappatura più sofisticata
                def _json_to_xml_element(data, root_name):
                    """Converte strutture JSON in elementi XML"""
                    if isinstance(data, dict):
                        root = ET.Element(root_name)
                        for key, value in data.items():
                            if key == "@attributes":
                                for attr_key, attr_val in value.items():
                                    root.set(attr_key, str(attr_val))
                            elif key == "#text":
                                root.text = str(value)
                            else:
                                root.append(_json_to_xml_element(value, key))
                        return root
                    elif isinstance(data, list):
                        root = ET.Element(root_name)
                        for item in data:
                            root.append(_json_to_xml_element(item, "item"))
                        return root
                    else:
                        # Caso base: elementi semplici
                        elem = ET.Element(root_name)
                        elem.text = str(data)
                        return elem
                
                # Parse JSON
                with open(input_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                # Converti in XML
                if isinstance(data, dict):
                    # Usa la prima chiave come nome dell'elemento root
                    root_name = next(iter(data))
                    root = _json_to_xml_element(data[root_name], root_name)
                else:
                    # Se è una lista o valore semplice, crea un elemento "root"
                    root = _json_to_xml_element(data, "root")
                
                # Salva come XML
                tree = ET.ElementTree(root)
                tree.write(output_path, encoding='utf-8', xml_declaration=True)
                
                return True
            
            elif source_format == 'csv' and target_format == 'xml':
                # Prima converti CSV -> JSON, poi JSON -> XML
                temp_json = tempfile.mktemp(suffix='.json')
                if self._convert_structured_formats(input_path, temp_json, 'csv', 'json'):
                    result = self._convert_structured_formats(temp_json, output_path, 'json', 'xml')
                    os.unlink(temp_json)
                    return result
                return False
            
            elif source_format == 'xml' and target_format == 'csv':
                # Prima converti XML -> JSON, poi JSON -> CSV
                temp_json = tempfile.mktemp(suffix='.json')
                if self._convert_structured_formats(input_path, temp_json, 'xml', 'json'):
                    result = self._convert_structured_formats(temp_json, output_path, 'json', 'csv')
                    os.unlink(temp_json)
                    return result
                return False
            
            else:
                print(f"Conversione da {source_format} a {target_format} non supportata per formati strutturati")
                return False
            
        except Exception as e:
            print(f"Errore nella conversione di formati strutturati: {e}")
            return False
    
    def _convert_with_pandoc(self, input_path, output_path, source_format, target_format, **kwargs):
        """Utilizza Pandoc per conversioni avanzate"""
        
        if not self.pandoc_available:
            print("Pandoc non disponibile per la conversione")
            return False
        
        try:
            # Costruisci i parametri pandoc
            cmd = ['pandoc', input_path, '-o', output_path]
            
            # Aggiungi opzioni in base ai parametri
            if kwargs.get('preserve_metadata', True):
                cmd.append('--standalone')
            
            if kwargs.get('paper_size'):
                cmd.extend(['-V', f'papersize={kwargs["paper_size"]}'])
            
            if kwargs.get('font'):
                cmd.extend(['-V', f'mainfont={kwargs["font"]}'])
            
            if kwargs.get('font_size'):
                cmd.extend(['-V', f'fontsize={kwargs["font_size"]}pt'])
            
            if kwargs.get('template'):
                cmd.extend(['--template', kwargs['template']])
            
            if kwargs.get('toc', False):
                cmd.append('--toc')
            
            # Gestisci i margini
            margins = kwargs.get('margins', {})
            for side in ['top', 'right', 'bottom', 'left']:
                if side in margins:
                    cmd.extend(['-V', f'margin-{side}={margins[side]}mm'])
            
            # Esegui pandoc
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode != 0:
                print(f"Errore Pandoc: {result.stderr}")
                return False
            
            return True
            
        except Exception as e:
            print(f"Errore nella conversione con Pandoc: {e}")
            return False
    
    def get_document_info(self, input_path):
        """
        Ottiene informazioni su un documento
        
        Args:
            input_path: percorso del documento
            
        Returns:
            dict: dizionario con le informazioni sul documento
        """
        try:
            ext = os.path.splitext(input_path)[1][1:].lower()
            info = {'format': ext}
            
            # Informazioni specifiche per tipo di file
            if ext == 'pdf':
                reader = PdfReader(input_path)
                info['pages'] = len(reader.pages)
                info['encrypted'] = reader.is_encrypted
                
                # Estrai metadati se disponibili
                if hasattr(reader.metadata, '_data'):
                    for key, value in reader.metadata._data.items():
                        if key.startswith('/'):
                            clean_key = key[1:]  # Rimuovi lo slash iniziale
                            info[clean_key] = value
                
            elif ext in ['docx', 'doc']:
                doc = docx.Document(input_path)
                
                # Numero di paragrafi e sezioni
                info['paragraphs'] = len(doc.paragraphs)
                info['sections'] = len(doc.sections)
                
                # Metadati del core properties
                if hasattr(doc, 'core_properties'):
                    props = doc.core_properties
                    for prop in ['author', 'category', 'comments', 'content_status', 
                                'created', 'identifier', 'keywords', 'language', 
                                'last_modified_by', 'last_printed', 'modified', 
                                'revision', 'subject', 'title', 'version']:
                        if hasattr(props, prop):
                            value = getattr(props, prop)
                            if value is not None:
                                info[prop] = str(value)
            
            # Informazioni generali del file
            file_stats = os.stat(input_path)
            info['size'] = file_stats.st_size
            info['created'] = file_stats.st_ctime
            info['modified'] = file_stats.st_mtime
            
            return info
                
        except Exception as e:
            print(f"Errore nel recupero delle informazioni del documento: {e}")
            return {'format': ext, 'error': str(e)}

    def _docx_to_pdf_fallback(self, input_path, output_path):
        """
        Fallback puro Python per convertire DOCX in PDF senza Word/LibreOffice.
        Utile per ambienti serverless come Render/Vercel.
        """
        try:
            doc = docx.Document(input_path)
            c = canvas.Canvas(output_path, pagesize=letter)
            width, height = letter
            
            # Margini
            margin_left = 50
            margin_top = 50
            y_position = height - margin_top
            
            # Font base
            c.setFont("Helvetica", 12)
            
            for para in doc.paragraphs:
                text = para.text
                if not text:
                    continue
                
                # Gestione semplice dello stile (grassetto per titoli)
                if para.style.name.startswith('Heading'):
                    c.setFont("Helvetica-Bold", 14)
                    y_position -= 10
                else:
                    c.setFont("Helvetica", 12)
                
                # Wrap del testo
                lines = textwrap.wrap(text, width=90)  # Approssimativo
                
                for line in lines:
                    if y_position < 50:  # Nuova pagina
                        c.showPage()
                        y_position = height - margin_top
                        c.setFont("Helvetica", 12)
                    
                    c.drawString(margin_left, y_position, line)
                    y_position -= 15
                
                y_position -= 10  # Spazio tra paragrafi
            
            c.save()
            return True
            
        except Exception as e:
            print(f"Errore nel fallback DOCX->PDF: {e}")
            return False